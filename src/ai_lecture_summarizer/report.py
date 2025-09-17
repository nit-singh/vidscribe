from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_title_page(doc: Document, title: str, subtitle: str | None = None) -> None:
	section = doc.sections[0]
	section.left_margin, section.right_margin = section.left_margin, section.right_margin
	title_par = doc.add_paragraph()
	title_run = title_par.add_run(title)
	title_run.bold = True
	title_run.font.size = Pt(28)
	title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if subtitle:
		sub_par = doc.add_paragraph()
		sub_run = sub_par.add_run(subtitle)
		sub_run.font.size = Pt(12)
		sub_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
	doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
	doc.add_heading(text, level=level)


def _add_bullets(doc: Document, lines: Iterable[str]) -> None:
	for line in lines:
		par = doc.add_paragraph(style="List Bullet")
		par.add_run(line)


def _parse_summary_to_doc(doc: Document, summary: str) -> None:
	lines = [l.rstrip() for l in summary.splitlines()]
	buffer: list[str] = []
	current_heading: str | None = None

	def flush_buffer_as_paragraph():
		if buffer:
			doc.add_paragraph(" ".join(buffer).strip())
			buffer.clear()

	for raw in lines:
		line = raw.strip()
		if not line:
			flush_buffer_as_paragraph()
			continue
		if line.startswith("## "):
			flush_buffer_as_paragraph()
			current_heading = line[3:].strip()
			_add_heading(doc, current_heading, level=2)
			continue
		if line.startswith("- ") or line.startswith("* "):
			_flush = " ".join(buffer).strip()
			if _flush:
				doc.add_paragraph(_flush)
				buffer.clear()
			_add_bullets(doc, [line[2:].strip()])
			continue
		buffer.append(line)

	flush_buffer_as_paragraph()


def build_docx_report(summary_text: str, output_path: Path, title: str, source_file: Path | None = None) -> Path:
	doc = Document()
	subtitle = None
	if source_file is not None:
		subtitle = f"Source: {source_file.name} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
	_add_title_page(doc, title=title, subtitle=subtitle)
	_add_heading(doc, "Executive Summary", level=1)
	_parse_summary_to_doc(doc, summary_text)
	output_path = Path(output_path)
	output_path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(str(output_path))
	return output_path
